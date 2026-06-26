"""
Agent layer for the Benefits Integrity Cloud FWA demo.

DEMO_MODE=true (default): every function returns pre-built fixtures immediately,
with zero Azure dependencies.

DEMO_MODE=false: replace the `# TODO(live)` stubs with real Azure AI Foundry
agent calls — one at a time, per the phased build-plan in docs/build-plan.md.
"""
import io
import json
import os
from functools import lru_cache
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


FOUNDRY_REASONING_MODEL = os.getenv("FOUNDRY_REASONING_MODEL", "gpt-4o")


def _live_config(name: str) -> str:
    value = os.getenv(name)
    if not value:
        raise RuntimeError(
            f"Live mode requires {name} to be set. "
            "Copy app/.env.example to .env and fill in the Foundry settings."
        )
    return value


@lru_cache(maxsize=1)
def _get_foundry_openai_client():
    try:
        from azure.ai.projects import AIProjectClient
        from azure.identity import DefaultAzureCredential
    except ImportError as exc:
        raise RuntimeError(
            "Live mode requires azure-ai-projects and azure-identity to be installed."
        ) from exc

    endpoint = _live_config("FOUNDRY_PROJECT_ENDPOINT")
    project_client = AIProjectClient(endpoint=endpoint, credential=DefaultAzureCredential())
    try:
        return project_client.get_openai_client()
    except Exception as exc:
        raise RuntimeError("Unable to create a Foundry OpenAI client for live mode.") from exc


def _run_foundry_json(system_prompt: str, payload: dict, model: str | None = None) -> dict:
    client = _get_foundry_openai_client()
    response = client.chat.completions.create(
        model=model or FOUNDRY_REASONING_MODEL,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False, indent=2)},
        ],
    )
    content = response.choices[0].message.content
    if not content:
        raise RuntimeError("Foundry returned an empty response.")
    parsed = json.loads(content)
    if not isinstance(parsed, dict):
        raise RuntimeError("Foundry returned a non-object JSON payload.")
    return parsed

# ---------------------------------------------------------------------------
# Scenario 1 — applicant intake
# ---------------------------------------------------------------------------


def run_intake_step(step: dict, demo: bool = True) -> dict:
    """Run (or simulate) a single intake check for Scenario 1.

    Returns the step enriched with an ``agent_note`` in demo mode, or calls
    a real Azure AI Foundry agent in live mode.
    """
    if demo:
        # Pass through the fixture data unchanged; add a demo watermark.
        return {
            "order": step["order"],
            "agent": step["agent"],
            "title": step["title"],
            "capabilities": step["capabilities"],
            "findings": step["result"]["findings"],
            "riskDelta": step["result"]["riskDelta"],
            "demo": True,
        }

    prompts = {
        "identity-proofing": (
            "You are the Identity Proofing Agent for a benefits intake flow. "
            "Return JSON only. Output the same step envelope the API uses, with "
            "keys: findings (array of strings), riskDelta (integer), "
            "credentialVerified (boolean), issuer (string or null), ok (boolean). "
            "Treat a missing trusted credential as a negative result. "
            "Use the provided input only and do not invent unrelated facts."
        ),
        "document-parsing+validation": (
            "You are the Document Parsing and Validation pipeline for a benefits "
            "intake flow. Return JSON only. Output the same step envelope the API "
            "uses, with keys: findings (array of strings), riskDelta (integer), "
            "tampered (boolean), expired (boolean), metadataEdited (string or null), "
            "regions (array of strings), docType (string or null), fields (object or null), "
            "confidence (number or null), ok (boolean). Apply fairness by routing "
            "low-confidence ambiguity to review rather than hard block."
        ),
        "validation": (
            "You are the Validation Agent for a benefits intake flow. Return JSON only. "
            "Output the same step envelope the API uses, with keys: findings (array of strings), "
            "riskDelta (integer), tampered (boolean), expired (boolean), metadataEdited "
            "(string or null), regions (array of strings), ok (boolean). "
            "Treat a single low-confidence blur as review, not block."
        ),
        "face-verification": (
            "You are the Face Verification Agent for a benefits intake flow. Return JSON only. "
            "Output the same step envelope the API uses, with keys: findings (array of strings), "
            "riskDelta (integer), live (boolean), deepfakeScore (number), faceMatch (boolean), ok (boolean). "
            "Use the provided video session information and ID portrait details only."
        ),
        "consistency-checker": (
            "You are the Consistency Checker for a benefits intake flow. Return JSON only. "
            "Output the same step envelope the API uses, with keys: findings (array of strings), "
            "riskDelta (integer), employerRegistered (boolean), incomeMatch (boolean), ok (boolean). "
            "Cross-check employer, income, and residency details against the supplied data."
        ),
    }
    prompt = prompts.get(
        step["agent"],
        "Return JSON only. Produce the same step envelope the API uses with findings, riskDelta, "
        "and any agent-specific fields inferred from the supplied data.",
    )
    payload = {k: v for k, v in step.items() if k != "result"}
    result = _run_foundry_json(prompt, payload)
    findings = result.get("findings", [])
    if not isinstance(findings, list):
        findings = [str(findings)]
    risk_delta = result.get("riskDelta")
    if risk_delta is None:
        raise RuntimeError("Live agent response did not include riskDelta.")
    output = {
        "order": step["order"],
        "agent": step["agent"],
        "title": step["title"],
        "capabilities": step["capabilities"],
        "findings": findings,
        "riskDelta": risk_delta,
        "ok": result.get("ok", risk_delta < 35),
        "demo": False,
    }
    for key, value in result.items():
        if key not in output:
            output[key] = value
    return output


def compose_decision(s1: dict, demo: bool = True) -> dict:
    """Compose the final eligibility decision for Scenario 1."""
    if demo:
        d = s1["decision"]
        return {
            "outcome": d["outcome"],
            "stage": d["stage"],
            "riskScore": d["riskScore"],
            "blockThreshold": d["blockThreshold"],
            "reasons": d["reasons"],
            "explainability": d["explainability"],
            "fairnessSafeguard": d["fairnessSafeguard"],
            "audit": d["audit"],
            "demo": True,
        }

    payload = {
        "scenario": s1.get("scenario", "01-prevention-blocked-pre-submission"),
        "steps": [
            {k: v for k, v in step.items() if k != "result"}
            for step in s1.get("steps", [])
        ],
        "policy": {
            "blockThreshold": 70,
            "reviewThreshold": 35,
            "riskBands": {
                "approve": "< 35",
                "review": "35..69",
                "block": ">= 70",
            },
        },
    }
    prompt = (
        "You are the Eligibility Decision Agent for a benefits intake flow. Return JSON only. "
        "Output the same decision envelope the API uses, with keys: outcome (approve|review|block), "
        "stage, riskScore (integer), blockThreshold (integer), reasons (array of strings), "
        "explainability (object with numeric weights), fairnessSafeguard (string), audit (object). "
        "Use the supplied steps and policy to derive the final decision. Keep the result explainable "
        "and consistent with the threshold policy."
    )
    result = _run_foundry_json(prompt, payload)
    outcome = result.get("outcome")
    risk_score = result.get("riskScore")
    if outcome is None or risk_score is None:
        raise RuntimeError("Live decision agent response was missing outcome or riskScore.")
    output = {
        "outcome": outcome,
        "stage": result.get("stage", "pre-submission"),
        "riskScore": risk_score,
        "blockThreshold": result.get("blockThreshold", 70),
        "reasons": result.get("reasons", []),
        "explainability": result.get("explainability", {}),
        "fairnessSafeguard": result.get("fairnessSafeguard", ""),
        "audit": result.get("audit", {"loggedTo": "Microsoft Purview"}),
        "demo": False,
    }
    return output


# ---------------------------------------------------------------------------
# Scenario 2 — officer / cluster detection
# ---------------------------------------------------------------------------


def build_graph(s2: dict, demo: bool = True) -> dict:
    """Return the relationship graph for a fraud cluster."""
    if demo:
        applications = s2["applications"]
        cluster = s2["cluster"]

        # Build node + edge lists from fixture data.
        nodes = []
        edges = []

        hub_meta = {
            "ip":    {"label": cluster["cyberSignals"]["sourceIp"], "color": "#3b9bff", "r": 17},
            "emp":   {"label": cluster["sharedEntities"]["employer"]["name"][:20], "color": "#9d7bff", "r": 15},
            "bank":  {"label": "IBAN ••4471", "color": "#37e0d8", "r": 15},
            "phone": {"label": "Phone ×2", "color": "#ffba49", "r": 13},
        }
        for hid, meta in hub_meta.items():
            nodes.append({"id": hid, "label": meta["label"], "type": hid,
                          "color": meta["color"], "r": meta["r"]})

        for i, app in enumerate(applications):
            ang = (3.14159 * 2 / len(applications)) * i - 3.14159 / 2
            import math
            x = 50 + math.cos(ang) * 40
            y = 50 + math.sin(ang) * 40
            nodes.append({"id": f"a{i}", "label": app["applicant"],
                          "type": "applicant", "color": "#ff5d76", "r": 6,
                          "x": round(x, 2), "y": round(y, 2)})
            # All applicants share IP and employer
            edges.append({"source": f"a{i}", "target": "ip"})
            edges.append({"source": f"a{i}", "target": "emp"})
            if app.get("sharedIban"):
                edges.append({"source": f"a{i}", "target": "bank"})
            if app.get("sharedPhone"):
                edges.append({"source": f"a{i}", "target": "phone"})

        return {"nodes": nodes, "edges": edges, "demo": True}

    payload = {
        "cluster": s2.get("cluster", {}),
        "applications": s2.get("applications", []),
        "graph": s2.get("graph", {}),
    }
    prompt = (
        "You are the Graph Relationship Explorer for a benefits integrity case. Return JSON only. "
        "Output a graph object with keys: nodes (array), edges (array), demo (boolean). Build the "
        "network from the supplied cluster and application data so shared IP, employer, bank, and phone "
        "relationships are explicit. Keep labels short and suitable for an investigator UI."
    )
    result = _run_foundry_json(prompt, payload)
    nodes = result.get("nodes")
    edges = result.get("edges")
    if nodes is None or edges is None:
        raise RuntimeError("Live graph agent response was missing nodes or edges.")
    return {
        "nodes": nodes,
        "edges": edges,
        "demo": False,
    }


def execute_action(action: str, s2: dict, demo: bool = True) -> dict:
    """Execute an officer action (freeze / escalate / casefile / notify)."""
    if demo:
        a = s2["actions"][action]
        audit_ts = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        return {
            "action": action,
            "status": a["status"],
            "label": a["label"],
            "agent": a["agent"],
            "detail": {k: v for k, v in a.items() if k not in ("label", "agent", "status")},
            "audit": {
                "timestamp": audit_ts,
                "loggedTo": "Microsoft Purview",
                "officer": "Maria S.",
                "clusterId": s2["cluster"]["id"],
            },
            "demo": True,
        }

    payload = {
        "action": action,
        "cluster": s2.get("cluster", {}),
        "actions": s2.get("actions", {}),
    }
    prompt = (
        "You are the officer action orchestrator for a benefits integrity case. Return JSON only. "
        "Output the same action envelope the API uses, with keys: action, status, label, agent, detail "
        "(object), audit (object), demo (boolean). Preserve governed, human-in-the-loop behavior and keep "
        "the response concise."
    )
    result = _run_foundry_json(prompt, payload)
    status = result.get("status")
    if status is None:
        raise RuntimeError("Live action agent response was missing status.")
    return {
        "action": result.get("action", action),
        "status": status,
        "label": result.get("label", s2["actions"][action]["label"]),
        "agent": result.get("agent", s2["actions"][action]["agent"]),
        "detail": result.get("detail", {}),
        "audit": result.get("audit", {}),
        "demo": False,
    }


# ---------------------------------------------------------------------------
# Case file generation (python-docx)
# ---------------------------------------------------------------------------

_EUR = lambda n: f"€{n:,.0f}"  # noqa: E731


def generate_case_file(s2: dict) -> bytes:
    """Generate a .docx fraud case file for a cluster and return the raw bytes."""
    cluster = s2["cluster"]
    applications = s2["applications"]
    totals = s2["totals"]
    signals = cluster["cyberSignals"]
    actions = s2["actions"]

    doc = Document()

    # ── styles ──────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    def heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4F, 0xA0)

    def add_kv(label: str, value: str, bad: bool = False):
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        run_v = p.add_run(value)
        if bad:
            run_v.font.color.rgb = RGBColor(0xFF, 0x3B, 0x5C)
        p.paragraph_format.space_after = Pt(2)

    # ── cover ────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"Fraud Case File — Cluster #{cluster['id']}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph(
        f"PREPARED BY MICROSOFT 365 COPILOT  ·  CONFIDENTIAL  ·  CASE ENF-{cluster['id']}"
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x9D, 0xB4, 0xE6)

    doc.add_paragraph(
        f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}  ·  "
        f"Risk score: {cluster['risk']} / 100  ·  CRITICAL"
    ).runs[0].bold = True

    doc.add_page_break()

    # ── 1. Executive summary ─────────────────────────────────────────────────
    heading("1. Executive summary")
    doc.add_paragraph(
        f"Nine benefit applications were submitted from a single IP address "
        f"({signals['sourceIp']}) within a {signals['submissionWindowSeconds']:.1f}-second window "
        f"by an automated agent, all citing the unregistered employer "
        f"\u201c{cluster['sharedEntities']['employer']['name']}\u201d and sharing a common bank account. "
        f"Cyber-signal, pattern and graph evidence indicate a coordinated organized-fraud ring. "
        f"All nine payments have been frozen and an enforcement case opened. "
        f"No improper funds were released."
    )

    # ── 2. Cyber-signal evidence ─────────────────────────────────────────────
    heading("2. Cyber-signal evidence")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Signal"
    hdr[1].text = "Finding"
    for hc in hdr:
        hc.paragraphs[0].runs[0].bold = True

    rows = [
        ("Source IP", signals["sourceIp"]),
        ("Network (ASN)", f"{signals['asnType'].title()} · {signals['asn']}"),
        ("Device fingerprint", f"Reused ×{signals['deviceFingerprintReuse']} (hash {signals['deviceFingerprint']})"),
        ("Form-fill cadence", f"~{signals['formFillMsPerField']} ms/field (bot-driven)"),
        ("Submission window", f"{signals['submissionWindowSeconds']:.1f} seconds ({cluster['applicationCount']} applications)"),
        ("Mean inter-arrival", f"{signals['meanInterArrivalSeconds']:.1f} seconds"),
        ("Threat-intel match", signals["threatIntelMatch"]),
    ]
    for label, val in rows:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = val

    # ── 3. Submission burst timeline ─────────────────────────────────────────
    heading("3. Submission burst timeline")
    doc.add_paragraph(
        f"Applications submitted {cluster['submissionTimestamps'][0]} to "
        f"{cluster['submissionTimestamps'][-1]} local time. "
        f"Mean inter-arrival {signals['meanInterArrivalSeconds']:.1f} seconds — "
        f"far below human entry speed and consistent with scripted automation."
    )

    # ── 4. Network relationships ──────────────────────────────────────────────
    heading("4. Network relationships")
    shared = cluster["sharedEntities"]
    for item in [
        f"Shared source IP across all {cluster['applicationCount']} applications",
        f"Shared employer: \u201c{shared['employer']['name']}\u201d (no active business / VAT registration)",
        f"Shared bank account / IBAN on {shared['ibanSharedOn']} of {cluster['applicationCount']} applications",
        f"Two shared contact phone numbers across the cluster",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ── 5. Application inventory ──────────────────────────────────────────────
    heading("5. Application inventory")
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(["Application ID", "Applicant", "Monthly benefit", "Status"]):
        hdr2[i].text = h
        hdr2[i].paragraphs[0].runs[0].bold = True

    for app in applications:
        r = table2.add_row().cells
        r[0].text = app["id"]
        r[1].text = app["applicant"]
        r[2].text = _EUR(app["monthlyBenefitEur"])
        r[3].text = "FROZEN"

    total_row = table2.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "Total monthly disbursement held"
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = _EUR(totals["monthlyBenefitHeldEur"])
    total_row[2].paragraphs[0].runs[0].bold = True
    total_row[3].text = "FROZEN"

    # ── 6. Responsible-AI risk explainability ─────────────────────────────────
    heading("6. Responsible-AI risk explainability")
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Table Grid"
    hdr3 = table3.rows[0].cells
    hdr3[0].text = "Contributing factor"
    hdr3[1].text = "Weight"
    for hc in hdr3:
        hc.paragraphs[0].runs[0].bold = True

    labels = {
        "coordinated_same_ip_burst": "Coordinated same-IP burst",
        "shared_fictitious_employer": "Shared fictitious employer",
        "shared_bank_account": "Shared bank account / IBAN",
        "bot_like_form_cadence": "Bot-like form cadence",
        "device_fingerprint_reuse": "Device fingerprint reuse",
    }
    for key, pct in cluster["riskExplainability"].items():
        r = table3.add_row().cells
        r[0].text = labels.get(key, key)
        r[1].text = f"{int(pct * 100)}%"

    # ── 7. Actions taken ──────────────────────────────────────────────────────
    heading("7. Actions taken")
    freeze_total = sum(a["monthlyBenefitEur"] for a in applications)
    for item in [
        f"Payments frozen — disbursement halted on all {cluster['applicationCount']} applications ({_EUR(freeze_total)}/mo held)",
        "Escalated to investigation — enforcement case ENF-4471, investigator J. Marković assigned, 48h SLA",
        "Case file generated — this document, by Document Generator + M365 Copilot",
        "Partner agencies notified — Tax Authority, Police Financial Crimes Unit, FIU (Purview-governed data-share)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ── 8. Recommended action ─────────────────────────────────────────────────
    heading("8. Recommended action")
    doc.add_paragraph(
        "Maintain the payment freeze, pursue enforcement against the identified ring, "
        "and recover any prior disbursements linked to the shared bank account. "
        "Reinstate individually any application later cleared through manual review."
    )

    # ── Governance note ───────────────────────────────────────────────────────
    note = doc.add_paragraph(
        "Fairness & governance: all automated findings are explainable and were confirmed "
        "by a human officer. Every decision and field-level disclosure is logged in "
        "Microsoft Purview."
    )
    note.runs[0].italic = True

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph(
        "\nDemonstration document · synthetic data. "
        "Names, companies, IP addresses and figures are fictitious and for illustration only."
    ).runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
